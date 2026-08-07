"""PDF / Word 文本提取脚本。

提取链路（PDF）：
    1. vision API（调同目录 vision.py 的 describe_image，逐页转图片识别）
    2. pdftotext 命令行工具
    3. pypdf / PyPDF2 纯 Python 解析

提取链路（Word）：
    1. python-docx
    2. unzip 解压 docx XML 提取文本

仅依赖 Python 标准库，所有可选依赖（vision / pdf2image / PyMuPDF(fitz) /
pypdf / PyPDF2 / python-docx）均使用 try import。
"""

from __future__ import annotations

import argparse
import os
import shutil
import subprocess
import sys
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from typing import List, Optional, Tuple

# 把本脚本所在目录加入 sys.path，确保能 import 同目录的 vision
_THIS_DIR = os.path.dirname(os.path.abspath(__file__))
if _THIS_DIR not in sys.path:
    sys.path.insert(0, _THIS_DIR)

# ---------- 可选依赖：try import ----------
try:
    import vision  # 同目录 vision.py
    _HAS_VISION = True
except Exception:
    vision = None
    _HAS_VISION = False

try:
    import fitz  # PyMuPDF
    _HAS_FITZ = True
except Exception:
    fitz = None
    _HAS_FITZ = False

try:
    from pdf2image import convert_from_path  # 需要 poppler
    _HAS_PDF2IMAGE = True
except Exception:
    convert_from_path = None
    _HAS_PDF2IMAGE = False

try:
    import pypdf
    _HAS_PYPDF = True
except Exception:
    pypdf = None
    _HAS_PYPDF = False

try:
    import PyPDF2
    _HAS_PYPDF2 = True
except Exception:
    PyPDF2 = None
    _HAS_PYPDF2 = False

try:
    import docx  # python-docx
    _HAS_DOCX = True
except Exception:
    docx = None
    _HAS_DOCX = False


# ---------- vision 可用性检查 ----------
def _vision_available() -> bool:
    """检查 vision API 是否可用（vision.py 可导入且配置了 API key）。"""
    if not _HAS_VISION:
        return False
    anthropic_key = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    openai_key = os.environ.get("OPENAI_API_KEY", "").strip()
    return bool(anthropic_key or openai_key)


def _pdftotext_available() -> bool:
    """检查 pdftotext 命令行工具是否可用。"""
    return shutil.which("pdftotext") is not None


# ---------- PDF 转图片 ----------
def _pdf_to_images(pdf_path: str) -> Tuple[List[str], Optional[str]]:
    """将 PDF 每页转为 PNG 图片。

    优先用 PyMuPDF(fitz)，其次 pdf2image（需 poppler）。
    返回 (图片路径列表, 临时目录路径)。若都不可用返回 ([], tmpdir)。
    """
    tmpdir = tempfile.mkdtemp(prefix="pdf_extract_")
    image_paths: List[str] = []

    # 1. PyMuPDF
    if _HAS_FITZ:
        try:
            doc = fitz.open(pdf_path)
            for i, page in enumerate(doc):
                pix = page.get_pixmap(dpi=200)
                out_path = os.path.join(tmpdir, f"page_{i:04d}.png")
                pix.save(out_path)
                image_paths.append(out_path)
            doc.close()
            return image_paths, tmpdir
        except Exception as exc:
            print(f"[pdf_extract] PyMuPDF 转图片失败: {exc}", file=sys.stderr)

    # 2. pdf2image
    if _HAS_PDF2IMAGE:
        try:
            images = convert_from_path(pdf_path, dpi=200)
            for i, img in enumerate(images):
                out_path = os.path.join(tmpdir, f"page_{i:04d}.png")
                img.save(out_path, "PNG")
                image_paths.append(out_path)
            return image_paths, tmpdir
        except Exception as exc:
            print(f"[pdf_extract] pdf2image 转图片失败: {exc}", file=sys.stderr)

    return image_paths, tmpdir


def _cleanup_tmpdir(tmpdir: Optional[str]) -> None:
    """清理临时目录。"""
    if not tmpdir:
        return
    try:
        shutil.rmtree(tmpdir, ignore_errors=True)
    except Exception:
        pass


# ---------- PDF 提取：各方式实现 ----------
def _extract_pdf_via_vision(pdf_path: str) -> str:
    """通过 vision API 逐页识别 PDF 内容，公式保持 LaTeX、表格用 markdown。"""
    image_paths, tmpdir = _pdf_to_images(pdf_path)
    try:
        if not image_paths:
            raise RuntimeError(
                "无法将 PDF 转为图片（需安装 PyMuPDF 或 pdf2image+poppler）"
            )
        prompt = (
            "请详细识别这一页的所有内容，"
            "数学公式用 LaTeX 格式输出，表格用 markdown 格式"
        )
        texts: List[str] = []
        total = len(image_paths)
        for i, img_path in enumerate(image_paths, start=1):
            print(
                f"[pdf_extract] vision 识别第 {i}/{total} 页...",
                file=sys.stderr,
            )
            text = vision.describe_image(img_path, prompt)
            texts.append(f"--- 第 {i} 页 ---\n{text}")
        return "\n\n".join(texts)
    finally:
        _cleanup_tmpdir(tmpdir)


def _extract_pdf_via_pdftotext(pdf_path: str) -> str:
    """用 pdftotext 命令行工具提取 PDF 文本。"""
    result = subprocess.run(
        ["pdftotext", "-layout", pdf_path, "-"],
        capture_output=True,
        text=True,
        timeout=600,
    )
    if result.returncode != 0:
        raise RuntimeError(
            f"pdftotext 返回码 {result.returncode}: {result.stderr.strip()}"
        )
    return result.stdout


def _extract_pdf_via_pypdf(pdf_path: str) -> str:
    """用 pypdf 或 PyPDF2 提取 PDF 文本。"""
    if _HAS_PYPDF:
        reader = pypdf.PdfReader(pdf_path)
        pages = reader.pages
    elif _HAS_PYPDF2:
        reader = PyPDF2.PdfReader(pdf_path)
        pages = reader.pages
    else:
        raise RuntimeError("pypdf / PyPDF2 均未安装")

    texts: List[str] = []
    for i, page in enumerate(pages, start=1):
        try:
            txt = page.extract_text() or ""
        except Exception as exc:
            print(f"[pdf_extract] 第 {i} 页提取失败: {exc}", file=sys.stderr)
            txt = ""
        texts.append(f"--- 第 {i} 页 ---\n{txt}")
    return "\n\n".join(texts)


# ---------- PDF 提取主入口 ----------
def extract_pdf(pdf_path: str, use_vision: bool = True) -> str:
    """提取 PDF 文本。

    优先级：vision API -> pdftotext -> pypdf/PyPDF2。
    前一种失败自动 fallback 到下一种。

    Args:
        pdf_path: PDF 文件路径。
        use_vision: 是否尝试 vision API（--no-vision 时设为 False）。

    Returns:
        提取到的文本。

    Raises:
        FileNotFoundError: 文件不存在。
        RuntimeError: 所有提取方式均失败。
    """
    if not os.path.isfile(pdf_path):
        raise FileNotFoundError(f"PDF 文件不存在: {pdf_path}")

    errors: List[str] = []

    # 1. vision API
    if use_vision and _vision_available():
        try:
            print("[pdf_extract] 尝试 vision API 提取...", file=sys.stderr)
            return _extract_pdf_via_vision(pdf_path)
        except Exception as exc:
            print(f"[pdf_extract] vision 提取失败: {exc}", file=sys.stderr)
            errors.append(f"vision: {exc}")
    elif use_vision and not _vision_available():
        print(
            "[pdf_extract] vision API 不可用（未配置 API key 或 vision.py 缺失），"
            "跳过 vision 提取",
            file=sys.stderr,
        )

    # 2. pdftotext
    if _pdftotext_available():
        try:
            print("[pdf_extract] 尝试 pdftotext 提取...", file=sys.stderr)
            return _extract_pdf_via_pdftotext(pdf_path)
        except Exception as exc:
            print(f"[pdf_extract] pdftotext 提取失败: {exc}", file=sys.stderr)
            errors.append(f"pdftotext: {exc}")
    else:
        print("[pdf_extract] pdftotext 不可用，跳过", file=sys.stderr)

    # 3. pypdf / PyPDF2
    if _HAS_PYPDF or _HAS_PYPDF2:
        try:
            print("[pdf_extract] 尝试 pypdf/PyPDF2 提取...", file=sys.stderr)
            return _extract_pdf_via_pypdf(pdf_path)
        except Exception as exc:
            print(f"[pdf_extract] pypdf 提取失败: {exc}", file=sys.stderr)
            errors.append(f"pypdf: {exc}")
    else:
        print("[pdf_extract] pypdf/PyPDF2 未安装，跳过", file=sys.stderr)

    raise RuntimeError(f"所有 PDF 提取方式均失败: {errors}")


# ---------- Word 提取 ----------
def _extract_docx_via_python_docx(docx_path: str) -> str:
    """用 python-docx 提取 docx 文本（含段落与表格）。"""
    doc = docx.Document(docx_path)
    parts: List[str] = []

    for para in doc.paragraphs:
        if para.text and para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells
            )
            if row_text.strip(" |"):
                parts.append(row_text)

    return "\n\n".join(parts)


def _extract_docx_via_unzip(docx_path: str) -> str:
    """解压 docx 的 word/document.xml 提取文本（fallback）。"""
    # docx XML 命名空间
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    t_tag = f"{{{w_ns}}}t"
    p_tag = f"{{{w_ns}}}p"

    texts: List[str] = []
    with zipfile.ZipFile(docx_path) as zf:
        try:
            with zf.open("word/document.xml") as f:
                tree = ET.parse(f)
        except KeyError:
            raise RuntimeError("docx 内未找到 word/document.xml（可能不是有效 docx）")
        root = tree.getroot()
        # 按段落 <w:p> 聚合 <w:t> 文本
        for para in root.iter(p_tag):
            para_text = "".join(
                t.text or "" for t in para.iter(t_tag)
            )
            if para_text.strip():
                texts.append(para_text)

    return "\n\n".join(texts)


def extract_word(docx_path: str) -> str:
    """提取 Word 文档文本。

    优先用 python-docx；失败 fallback 到 unzip XML 提取。
    注意：旧版 .doc（二进制格式）python-docx 不支持，unzip 方式也仅适用于
    .docx（zip 容器）。.doc 二进制会在此处抛出可读错误。

    Args:
        docx_path: Word 文件路径 (.docx / .doc)。

    Returns:
        提取到的文本。

    Raises:
        FileNotFoundError: 文件不存在。
        RuntimeError: 所有提取方式均失败。
    """
    if not os.path.isfile(docx_path):
        raise FileNotFoundError(f"Word 文件不存在: {docx_path}")

    errors: List[str] = []

    if _HAS_DOCX:
        try:
            print("[pdf_extract] 尝试 python-docx 提取...", file=sys.stderr)
            return _extract_docx_via_python_docx(docx_path)
        except Exception as exc:
            print(f"[pdf_extract] python-docx 提取失败: {exc}", file=sys.stderr)
            errors.append(f"python-docx: {exc}")
    else:
        print("[pdf_extract] python-docx 未安装，尝试 unzip 方式", file=sys.stderr)

    try:
        print("[pdf_extract] 尝试 unzip XML 提取...", file=sys.stderr)
        return _extract_docx_via_unzip(docx_path)
    except Exception as exc:
        print(f"[pdf_extract] unzip 提取失败: {exc}", file=sys.stderr)
        errors.append(f"unzip: {exc}")

    ext = os.path.splitext(docx_path)[1].lower()
    hint = ""
    if ext == ".doc":
        hint = (
            "；提示：旧版 .doc 为二进制格式，python-docx 与 unzip 均不支持，"
            "建议先转换为 .docx"
        )
    raise RuntimeError(f"Word 文档提取失败: {errors}{hint}")


# ---------- 统一分发入口 ----------
def extract(file_path: str) -> str:
    """根据扩展名分发到 extract_pdf 或 extract_word。

    支持 .pdf / .docx / .doc。

    Args:
        file_path: 文件路径。

    Returns:
        提取到的文本。

    Raises:
        FileNotFoundError: 文件不存在。
        ValueError: 不支持的文件格式。
        RuntimeError: 提取失败。
    """
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_pdf(file_path, use_vision=True)
    elif ext in (".docx", ".doc"):
        return extract_word(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


# ---------- CLI ----------
def _build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        prog="pdf_extract.py",
        description="PDF/Word 文本提取工具（vision API / pdftotext / pypdf / python-docx）。",
    )
    parser.add_argument("file_path", help="要提取的文件路径 (.pdf/.docx/.doc)")
    parser.add_argument(
        "--out",
        default=None,
        help="同时写入到指定文件（默认仅输出到 stdout）",
    )
    parser.add_argument(
        "--no-vision",
        action="store_true",
        help="强制不使用 vision API（PDF 走 pdftotext / pypdf）",
    )
    return parser


def main(argv: Optional[list] = None) -> int:
    parser = _build_arg_parser()
    args = parser.parse_args(argv)

    ext = os.path.splitext(args.file_path)[1].lower()
    try:
        if ext == ".pdf":
            text = extract_pdf(args.file_path, use_vision=not args.no_vision)
        elif ext in (".docx", ".doc"):
            text = extract_word(args.file_path)
        else:
            print(f"[pdf_extract] 不支持的文件格式: {ext}", file=sys.stderr)
            return 1
    except Exception as exc:
        print(f"[pdf_extract] 错误: {exc}", file=sys.stderr)
        return 1

    # 输出到 stdout
    print(text)

    # 可选写入文件
    if args.out:
        try:
            with open(args.out, "w", encoding="utf-8") as f:
                f.write(text)
            print(f"[pdf_extract] 已写入: {args.out}", file=sys.stderr)
        except Exception as exc:
            print(f"[pdf_extract] 写入文件失败: {exc}", file=sys.stderr)
            return 1

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
